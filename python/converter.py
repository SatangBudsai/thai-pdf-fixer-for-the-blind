#!/usr/bin/env python3
"""
Thai PDF to accessible DOCX converter.
Sidecar process for the Tauri desktop app.

Usage:
  converter convert <input.pdf> <output.docx>   — full PDF-to-Word conversion
  converter preview <input.pdf>                  — extract text only (for TTS)

Outputs JSON lines to stdout for progress reporting.
"""

import sys
import json
import os
import io

# Force UTF-8 stdout on Windows (default cp1252 can't handle Thai)
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")

import re
import fitz  # PyMuPDF
import pdfplumber
from pythainlp.util import normalize as thai_normalize
from docx import Document
from docx.shared import Pt, Inches


# ════════════════════════════════════════════════════════════════════
# Thai text fixers
# ════════════════════════════════════════════════════════════════════

def fix_sara_am(text: str) -> str:
    """
    Fix decomposed sara am (ำ) that PDFs often break into:
      - consonant + space(s) + sara aa (า)  →  consonant + sara am (ำ)
      - nikhahit (◌ํ) + sara aa (า)          →  sara am (ำ)
    Thai consonants: U+0E01–U+0E2E
    """
    # Fix: nikhahit (U+0E4D) + sara aa (U+0E32) → sara am (U+0E33)
    text = text.replace('\u0E4D\u0E32', '\u0E33')
    # Fix: consonant + whitespace + sara aa → consonant + sara am
    text = re.sub(r'([\u0E01-\u0E2E])\s+\u0E32', '\\1\u0E33', text)
    return text


def normalize_thai(text: str) -> str:
    """Apply all Thai text normalization steps."""
    text = fix_sara_am(text)
    text = thai_normalize(text)
    return text


# ════════════════════════════════════════════════════════════════════
# JSON-line protocol helpers
# ════════════════════════════════════════════════════════════════════

def emit(obj: dict):
    """Write a JSON line to stdout and flush immediately."""
    print(json.dumps(obj, ensure_ascii=False), flush=True)


def emit_progress(page: int, total: int, message: str):
    emit({"type": "progress", "page": page, "total": total, "message": message})


def emit_error(message: str):
    emit({"type": "error", "message": message})


# ════════════════════════════════════════════════════════════════════
# Geometry helper
# ════════════════════════════════════════════════════════════════════

def rects_overlap(r1, r2):
    """
    Check if two rectangles overlap.
    Each rect is (x0, y0, x1, y1) where (x0,y0) is top-left, (x1,y1) is bottom-right.
    """
    return not (r1[2] <= r2[0] or r1[0] >= r2[2] or r1[3] <= r2[1] or r1[1] >= r2[3])


# ════════════════════════════════════════════════════════════════════
# Page extraction
# ════════════════════════════════════════════════════════════════════

def extract_page_content(plumber_page, fitz_page):
    """
    Extract text, tables, and images from a single PDF page.

    Strategy:
    1. Use pdfplumber to find tables and their bounding boxes.
    2. Use PyMuPDF to get text blocks, filtering out any that overlap
       with table regions (so text inside tables isn't duplicated).
    3. Use pdfplumber to extract table data as 2D arrays.
    4. Use PyMuPDF to extract embedded images.

    Returns a list of content items sorted by vertical position:
      {"type": "text",  "content": "...", "y": float}
      {"type": "table", "content": [["cell", ...], ...], "y": float}
      {"type": "image", "data": bytes, "ext": str, "y": float}
    """
    content_items = []

    # ── Step 1: Detect tables ──────────────────────────────────────
    tables = plumber_page.find_tables()
    table_bboxes = [t.bbox for t in tables]  # (x0, y0, x1, y1)

    # ── Step 2: Extract text blocks outside table regions ──────────
    # PyMuPDF get_text("blocks") returns:
    #   (x0, y0, x1, y1, text_or_image, block_no, block_type)
    #   block_type: 0 = text, 1 = image
    text_blocks = fitz_page.get_text("blocks")

    for block in text_blocks:
        block_type = block[6]
        if block_type != 0:  # skip image blocks (handled separately)
            continue

        bx0, by0, bx1, by1 = block[:4]
        text = block[4].strip()

        if not text:
            continue

        # Skip if this text block overlaps any table bounding box
        in_table = any(rects_overlap((bx0, by0, bx1, by1), tb) for tb in table_bboxes)
        if in_table:
            continue

        # Apply Thai text normalization (fix ำ decomposition + PyThaiNLP)
        normalized = normalize_thai(text)
        content_items.append({"type": "text", "content": normalized, "y": by0})

    # ── Step 3: Extract table data ─────────────────────────────────
    for table in tables:
        table_data = table.extract()
        if not table_data:
            continue

        # Normalize Thai text in each cell
        normalized_table = []
        for row in table_data:
            normalized_row = []
            for cell in row:
                if cell:
                    normalized_row.append(normalize_thai(cell))
                else:
                    normalized_row.append("")
            normalized_table.append(normalized_row)

        # Use the table's top Y coordinate for ordering
        content_items.append({
            "type": "table",
            "content": normalized_table,
            "y": table.bbox[1]
        })

    # ── Step 4: Extract images ─────────────────────────────────────
    images = fitz_page.get_images(full=True)
    for img_info in images:
        xref = img_info[0]
        try:
            base_image = fitz_page.parent.extract_image(xref)
            if not base_image or not base_image.get("image"):
                continue

            img_bytes = base_image["image"]
            img_ext = base_image.get("ext", "png")

            # Skip very small images (likely icons/artifacts, < 2KB)
            if len(img_bytes) < 2048:
                continue

            # Always convert through PIL to fix DPI=0 and unsupported formats
            try:
                from PIL import Image as PILImage
                pil_img = PILImage.open(io.BytesIO(img_bytes))
                buf = io.BytesIO()
                # Convert RGBA/P to RGB for JPEG compatibility, save as PNG
                if pil_img.mode in ("RGBA", "P", "LA"):
                    pil_img = pil_img.convert("RGBA")
                buf = io.BytesIO()
                pil_img.save(buf, format="PNG")
                img_bytes = buf.getvalue()
                img_ext = "png"
            except Exception:
                continue

            # Determine the position and display size of the image on the page
            img_rects = fitz_page.get_image_rects(xref)
            if img_rects:
                rect = img_rects[0]
                y_pos = rect.y0
                width_pts = rect.width
                height_pts = rect.height
            else:
                y_pos = 0.0
                width_pts = 0
                height_pts = 0

            content_items.append({
                "type": "image",
                "data": img_bytes,
                "ext": img_ext,
                "y": y_pos,
                "width_pts": width_pts,
                "height_pts": height_pts,
            })
        except Exception:
            continue

    # Sort all items by vertical position so they appear in reading order
    content_items.sort(key=lambda item: item["y"])

    return content_items


# ════════════════════════════════════════════════════════════════════
# DOCX builder
# ════════════════════════════════════════════════════════════════════

def build_docx(all_pages_content, output_path):
    """
    Build a Word document from extracted page content.
    Each page's content is a list of items (text, table, image).
    """
    doc = Document()

    # Set default font to TH SarabunPSK for Thai readability
    style = doc.styles["Normal"]
    font = style.font
    font.name = "TH SarabunPSK"
    font.size = Pt(16)

    for page_idx, page_content in enumerate(all_pages_content):
        if page_idx > 0:
            doc.add_page_break()

        for item in page_content:
            if item["type"] == "text":
                doc.add_paragraph(item["content"])

            elif item["type"] == "table":
                table_data = item["content"]
                if not table_data:
                    continue
                num_cols = max(len(row) for row in table_data)
                tbl = doc.add_table(rows=len(table_data), cols=num_cols)
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(table_data):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < num_cols:
                            cell = tbl.rows[r_idx].cells[c_idx]
                            cell.text = cell_text or ""
                            # Apply Thai font to table cells too
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = "TH SarabunPSK"
                                    run.font.size = Pt(14)

            elif item["type"] == "image":
                try:
                    image_stream = io.BytesIO(item["data"])
                    # Use proportional width from PDF, clamped to 1–6.5 inches
                    width_pts = item.get("width_pts", 0)
                    if width_pts > 0:
                        width_in = max(1.0, min(6.5, width_pts / 72.0))
                    else:
                        width_in = 5.0
                    doc.add_picture(image_stream, width=Inches(width_in))
                except Exception:
                    # If image insertion fails, add a placeholder
                    doc.add_paragraph("[ไม่สามารถแทรกรูปภาพได้]")

    doc.save(output_path)


# ════════════════════════════════════════════════════════════════════
# Commands
# ════════════════════════════════════════════════════════════════════

def cmd_convert(input_path: str, output_path: str):
    """Full PDF → DOCX conversion."""
    fitz_doc = fitz.open(input_path)
    plumber_doc = pdfplumber.open(input_path)
    total_pages = len(fitz_doc)

    all_pages_content = []

    for page_num in range(total_pages):
        emit_progress(
            page_num + 1, total_pages,
            f"กำลังประมวลผลหน้า {page_num + 1} จาก {total_pages}"
        )

        fitz_page = fitz_doc[page_num]
        plumber_page = plumber_doc.pages[page_num]
        page_content = extract_page_content(plumber_page, fitz_page)
        all_pages_content.append(page_content)

    emit_progress(total_pages, total_pages, "กำลังสร้างไฟล์ Word...")
    build_docx(all_pages_content, output_path)

    fitz_doc.close()
    plumber_doc.close()

    emit({"type": "done", "output": output_path, "pages": total_pages})


def cmd_preview(input_path: str):
    """Extract and normalize text only — output for TTS preview."""
    fitz_doc = fitz.open(input_path)
    plumber_doc = pdfplumber.open(input_path)
    total_pages = len(fitz_doc)

    for page_num in range(total_pages):
        emit_progress(
            page_num + 1, total_pages,
            f"กำลังอ่านหน้า {page_num + 1} จาก {total_pages}"
        )

        fitz_page = fitz_doc[page_num]
        plumber_page = plumber_doc.pages[page_num]
        page_content = extract_page_content(plumber_page, fitz_page)

        # Collect all text from this page (text blocks + table cells)
        page_texts = []
        for item in page_content:
            if item["type"] == "text":
                page_texts.append(item["content"])
            elif item["type"] == "table":
                for row in item["content"]:
                    row_text = " | ".join(cell for cell in row if cell)
                    if row_text:
                        page_texts.append(row_text)

        if page_texts:
            emit({
                "type": "text",
                "page": page_num + 1,
                "content": "\n".join(page_texts)
            })

    fitz_doc.close()
    plumber_doc.close()

    emit({"type": "preview_done", "pages": total_pages})


# ════════════════════════════════════════════════════════════════════
# Main entry point
# ════════════════════════════════════════════════════════════════════

def main():
    if len(sys.argv) < 2:
        emit_error("Usage: converter <convert|preview> <input.pdf> [output.docx]")
        sys.exit(1)

    command = sys.argv[1]

    try:
        if command == "convert":
            if len(sys.argv) != 4:
                emit_error("Usage: converter convert <input.pdf> <output.docx>")
                sys.exit(1)

            input_path = sys.argv[2]
            output_path = sys.argv[3]

            if not os.path.exists(input_path):
                emit_error(f"ไม่พบไฟล์: {input_path}")
                sys.exit(1)

            cmd_convert(input_path, output_path)

        elif command == "preview":
            if len(sys.argv) != 3:
                emit_error("Usage: converter preview <input.pdf>")
                sys.exit(1)

            input_path = sys.argv[2]

            if not os.path.exists(input_path):
                emit_error(f"ไม่พบไฟล์: {input_path}")
                sys.exit(1)

            cmd_preview(input_path)

        else:
            emit_error(f"Unknown command: {command}. Use 'convert' or 'preview'.")
            sys.exit(1)

    except PermissionError:
        emit({"type": "error", "code": "file_locked", "message": "ไม่สามารถบันทึกไฟล์ได้ กรุณาปิดไฟล์ Word ก่อนแล้วลองใหม่"})
        sys.exit(1)
    except Exception as e:
        emit_error(f"เกิดข้อผิดพลาด: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()
